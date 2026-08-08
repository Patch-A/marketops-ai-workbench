#!/usr/bin/env python3
"""Generate the synthetic DOCX used by the M0-02 parser spike."""

from __future__ import annotations

import argparse
import os
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


CONTENT_WIDTH_DXA = 9360
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
REPRODUCIBLE_ZIP_TIME = (2024, 1, 1, 0, 0, 0)


def set_run_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style(style, *, size, color, before, after, line_spacing=1.10, bold=None):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing


def ensure_child(parent, tag):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_width(parent, tag, width):
    node = ensure_child(parent, tag)
    node.set(qn("w:type"), "dxa")
    node.set(qn("w:w"), str(width))


def apply_table_geometry(table, widths):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("column widths must total 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    set_width(table_pr, "w:tblW", CONTENT_WIDTH_DXA)
    indent = ensure_child(table_pr, "w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    layout = ensure_child(table_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        row.height = None
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            properties = cell._tc.get_or_add_tcPr()
            set_width(properties, "w:tcW", widths[index])
            margins = ensure_child(properties, "w:tcMar")
            for side, value in CELL_MARGINS.items():
                margin = ensure_child(margins, f"w:{side}")
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")


def mark_header_row(row):
    row_properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    row_properties.append(header)


def shade_cell(cell, fill):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def fill_table(table, headers, rows, widths):
    for index, value in enumerate(headers):
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(paragraph.add_run(value), bold=True)
        shade_cell(table.rows[0].cells[index], "F2F4F7")
    mark_header_row(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            set_run_font(paragraph.add_run(value))
    apply_table_geometry(table, widths)


def normalize_docx_package(path: Path):
    """Rewrite the OOXML ZIP with stable order, timestamps, and metadata."""
    with ZipFile(path, "r") as source:
        entries = [(name, source.read(name), source.getinfo(name)) for name in sorted(source.namelist())]
    handle, temporary_name = tempfile.mkstemp(suffix=".docx", dir=path.parent)
    os.close(handle)
    temporary = Path(temporary_name)
    try:
        with ZipFile(temporary, "w", compression=ZIP_DEFLATED, compresslevel=9) as target:
            for name, payload, original in entries:
                info = ZipInfo(name, date_time=REPRODUCIBLE_ZIP_TIME)
                info.compress_type = ZIP_DEFLATED
                info.external_attr = original.external_attr
                info.create_system = original.create_system
                target.writestr(info, payload)
        os.replace(temporary, path)
    finally:
        temporary.unlink(missing_ok=True)


def build(output: Path):
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style(document.styles["Normal"], size=11, color="000000", before=0, after=6)
    set_style(document.styles["Heading 1"], size=16, color="2E74B5", before=16, after=8, bold=True)
    set_style(document.styles["Heading 2"], size=13, color="2E74B5", before=12, after=6, bold=True)
    set_style(document.styles["Heading 3"], size=12, color="1F4D78", before=8, after=4, bold=True)

    # standard_business_brief with a named fixture-title override.
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("凌川工业增长峰会 AI 工作简报"), size=23, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle.add_run("SYNTHETIC FIXTURE | M0-02 document parsing"), size=10, color="555555", bold=True)

    document.add_heading("1. 已确认项目输入", level=1)
    document.add_paragraph("活动日为 2026-10-23，形式为单日线下 B2B 峰会。所有品牌、人物和项目数据均为合成内容。")
    document.add_heading("1.1 活动目标", level=2)
    document.add_paragraph("主办方确认内容触达、重点客户会面和可追溯复盘是三个项目目标。")

    source_table = document.add_table(rows=1, cols=3)
    source_table.style = "Table Grid"
    fill_table(
        source_table,
        ["编号", "已确认事实", "来源"],
        [
            ["SRC-DOCX-01", "场地容量上限为 360 人。", "合成 Brief"],
            ["SRC-DOCX-02", "最终进场清单必须在 2026-10-16 前批准。", "合成场地规则"],
            ["SRC-DOCX-03", "正式排期变更需要项目负责人批准。", "合成审批规则"],
        ],
        [1600, 5360, 2400],
    )

    document.add_heading("2. 执行约束", level=1)
    document.add_paragraph("预算、供应商报价和 ROI 均为未知；解析结果不得补写缺失数字。")
    document.add_paragraph("原始文件哈希与来源坐标必须随解析块保存，文件变化后需要重新解析。")

    document.add_heading("3. 变更审批", level=1)
    document.add_paragraph("下表用于验证表格顺序、单元格坐标和章节路径。")
    change_table = document.add_table(rows=1, cols=4)
    change_table.style = "Table Grid"
    fill_table(
        change_table,
        ["变更编号", "影响对象", "状态", "审批角色"],
        [
            ["CHG-DOCX-01", "DEL-03 / TASK-080", "待审批", "项目负责人"],
            ["CHG-DOCX-02", "无正式影响", "已驳回", "市场负责人"],
        ],
        [1800, 3000, 1800, 2760],
    )

    document.core_properties.title = "M0-02 Synthetic Document Parser Fixture"
    document.core_properties.subject = "Synthetic validation data"
    document.core_properties.author = "MarketOps AI Workbench"
    document.core_properties.keywords = "synthetic, parser, validation"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    normalize_docx_package(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.output)
    print(f"Generated {args.output}")


if __name__ == "__main__":
    main()
