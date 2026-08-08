from pathlib import Path
from zipfile import ZipFile
import re

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AI市场工作台_二手证据与市场验证执行包_审阅版.docx"


def main():
    doc = Document(DOCX)
    section = doc.sections[0]
    assertions = {
        "page_size_letter": round(section.page_width.inches, 2) == 8.50
        and round(section.page_height.inches, 2) == 11.00,
        "margins_1in": all(
            round(x.inches, 2) == 1.00
            for x in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin)
        ),
        "header_footer_distance": round(section.header_distance.inches, 3) == 0.492
        and round(section.footer_distance.inches, 3) == 0.492,
        "substantial_content": len(doc.paragraphs) > 250,
        "expected_tables": len(doc.tables) >= 8,
    }

    all_text = "\n".join(p.text for p in doc.paragraphs)
    all_text += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assertions["parts_present"] = "Part I  二手证据与产品判断" in all_text and "Part II  市场验证执行包" in all_text
    assertions["decision_present"] = "Conditional Go" in all_text
    assertions["no_placeholders"] = not re.search(r"TODO|TBD|lorem ipsum|\{\{|\}\}|某某", all_text, re.I)
    assertions["no_mojibake"] = not any(token in all_text for token in ("锛", "鈥", "銆", "绔", "馃"))

    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and tbl_w.get(qn("w:w")) == "9360"
        grid = [int(x.get(qn("w:w"))) for x in table._tbl.tblGrid]
        assert sum(grid) == 9360
        for row in table.rows:
            assert len(row.cells) == len(grid)

    with ZipFile(DOCX) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
        assertions["real_numbering"] = "<w:numPr>" in document_xml and "w:numFmt w:val=\"bullet\"" in numbering_xml
        assertions["preset_styles"] = all(x in styles_xml for x in ("2E74B5", "1F4D78", "Microsoft YaHei"))
        assertions["hyperlinks"] = rels_xml.count("relationships/hyperlink") >= 10
        assertions["page_field"] = " PAGE " in archive.read("word/footer1.xml").decode("utf-8")

    failures = [name for name, ok in assertions.items() if not ok]
    for name, ok in assertions.items():
        print(f"{'PASS' if ok else 'FAIL'}\t{name}")
    print(f"paragraphs\t{len(doc.paragraphs)}")
    print(f"tables\t{len(doc.tables)}")
    print(f"characters\t{len(all_text)}")
    if failures:
        raise SystemExit("Failed checks: " + ", ".join(failures))


if __name__ == "__main__":
    main()
