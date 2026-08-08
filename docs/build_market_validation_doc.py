from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
EVIDENCE_MD = ROOT / "docs" / "market-validation-evidence.md"
PLAYBOOK_MD = ROOT / "docs" / "market-validation-playbook.md"
OUTPUT = ROOT / "AI市场工作台_二手证据与市场验证执行包_审阅版.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "202124"
MUTED = "5F6368"
LIGHT_GRAY = "F2F4F7"
LINE = "D5D9E0"
WHITE = "FFFFFF"
CAUTION = "7A5A00"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, size=None, bold=None, italic=None, color=INK, code=False):
    latin = "Consolas" if code else "Calibri"
    east_asia = "Microsoft YaHei"
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = rgb(color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=LINE, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def configure_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def hyperlink(paragraph, label, url, bold=False):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.extend([fonts, color, underline])
    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    link.append(run)
    paragraph._p.append(link)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\(https?://[^)]+\))")


def add_inline(paragraph, text, base_size=11, base_color=INK):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True, color=base_color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size - 0.5, color=DARK_BLUE, code=True)
        else:
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            hyperlink(paragraph, label, url)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=base_color)


def add_numbering_definitions(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(abstract_ids or [0]) + 1
    next_num = max(num_ids or [0]) + 1
    result = {}

    for kind, fmt, text_value, font in (
        ("bullet", "bullet", "•", "Symbol"),
        ("decimal", "decimal", "%1.", "Calibri"),
    ):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(next_abs))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, ind, spacing])
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), font)
        fonts.set(qn("w:hAnsi"), font)
        fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r_pr.append(fonts)
        lvl.extend([start, num_fmt, lvl_text, jc, p_pr, r_pr])
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(next_abs))
        num.append(abs_id)
        numbering.append(num)
        result[kind] = next_num
        next_abs += 1
        next_num += 1
    return result


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if value and node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    elif not value and node is not None:
        p_pr.remove(node)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("Table Text", "Table Header", "Metadata", "Source Note", "Callout"):
        if style_name not in styles:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    table_text = styles["Table Text"]
    table_text.font.name = "Calibri"
    table_text.font.size = Pt(9)
    table_text._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(2)
    table_text.paragraph_format.line_spacing = 1.0

    table_header = styles["Table Header"]
    table_header.font.name = "Calibri"
    table_header.font.size = Pt(9)
    table_header.font.bold = True
    table_header.font.color.rgb = rgb(INK)
    table_header._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    table_header.paragraph_format.space_before = Pt(0)
    table_header.paragraph_format.space_after = Pt(2)
    table_header.paragraph_format.line_spacing = 1.0

    metadata = styles["Metadata"]
    metadata.font.name = "Calibri"
    metadata.font.size = Pt(10)
    metadata.font.color.rgb = rgb(MUTED)
    metadata._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(2)
    metadata.paragraph_format.line_spacing = 1.0

    source_note = styles["Source Note"]
    source_note.font.name = "Calibri"
    source_note.font.size = Pt(9)
    source_note.font.color.rgb = rgb(MUTED)
    source_note._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    source_note.paragraph_format.space_before = Pt(4)
    source_note.paragraph_format.space_after = Pt(4)
    source_note.paragraph_format.line_spacing = 1.0

    callout = styles["Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(11)
    callout.font.bold = True
    callout.font.color.rgb = rgb(DARK_BLUE)
    callout._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.10


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("AI 市场工作台 | 证据验证与市场测试")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("审阅版  |  ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(p)
    return section


def add_masthead(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("研究与验证执行包")
    set_run_font(run, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("AI 市场工作台")
    set_run_font(run, size=25, bold=True, color="000000")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("二手证据验证、MVP 判断与真实 brief 市场测试")
    set_run_font(run, size=14, color="373737")

    rows = [
        ("状态", "Conditional Go：进入真实任务验证，不直接开发完整 Agent"),
        ("适用对象", "项目发起人、活动/品牌代理公司、B2B 市场验证团队"),
        ("版本", "v0.1 审阅版"),
        ("日期", "2026-08-08"),
        ("证据截止", "2026-08-08；公开来源均保留链接、样本和局限"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    configure_table_geometry(table, [1500, 7860])
    set_table_borders(table, color=WHITE, size="0")
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        left.paragraphs[0].style = "Metadata"
        right.paragraphs[0].style = "Metadata"
        left.paragraphs[0].clear()
        right.paragraphs[0].clear()
        r = left.paragraphs[0].add_run(label)
        set_run_font(r, size=10, bold=True, color="000000")
        add_inline(right.paragraphs[0], value, base_size=10, base_color="000000")

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(12)
    rule.paragraph_format.space_after = Pt(14)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    p = doc.add_paragraph(style="Callout")
    add_inline(
        p,
        "结论：公开证据支持验证一个窄而可控的证据化方案工作流；尚未证明用户会为完整软件付费。",
        base_size=11,
        base_color=DARK_BLUE,
    )

    p = doc.add_paragraph()
    add_inline(
        p,
        "建议先卖真实 brief 的交付结果，再决定是否把验证过的流程产品化。Daily AI Radar 用于后续学习与留存，不是首个商业切入点。",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("审阅重点")
    set_run_font(r, size=11, bold=True, color="000000")
    for item in (
        "是否同意以活动/品牌代理公司为主组、B2B 企业市场为对照组？",
        "是否能在两周内取得 3 个历史项目和 1 个实时脱敏 brief？",
        "测试阈值和人民币 299/699 元的付费锚点是否需要调整？",
    ):
        p = doc.add_paragraph()
        apply_numbering(p, doc._numbering_ids["bullet"])
        add_inline(p, item)

    doc.add_page_break()


def add_contents(doc):
    p = doc.add_paragraph("文档结构", style="Heading 1")
    set_keep_with_next(p)
    for item in (
        "Part I  二手证据、机会判断与 MVP 边界",
        "Part II  市场验证执行包：招募、回放、评分、付费测试",
        "使用说明  所有百分比必须区分外部证据、测试阈值和真实结果",
    ):
        p = doc.add_paragraph()
        apply_numbering(p, doc._numbering_ids["bullet"])
        add_inline(p, item)
    p = doc.add_paragraph(style="Source Note")
    add_inline(
        p,
        "说明：本文件是审阅和执行材料，不是融资报告或市场规模证明。网络研究不能替代真实行为验证。",
        base_size=9,
        base_color=MUTED,
    )
    doc.add_page_break()


def table_widths(headers):
    n = len(headers)
    if n == 2:
        return [2700, 6660]
    if n == 3:
        return [1700, 3180, 4480]
    if n == 4:
        if any("权重" in h or "分" in h or "样本" in h for h in headers):
            return [1900, 900, 3300, 3260]
        return [1500, 1800, 3000, 3060]
    if n == 5:
        return [1300, 1850, 1850, 2180, 2180]
    return [9360 // n] * n


def add_markdown_table(doc, rows):
    headers = [x.strip() for x in rows[0].strip().strip("|").split("|")]
    body = rows[2:]
    table = doc.add_table(rows=1 + len(body), cols=len(headers))
    widths = table_widths(headers)
    widths[-1] += 9360 - sum(widths)
    configure_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.style = "Table Header"
        p.clear()
        add_inline(p, text, base_size=9, base_color=INK)

    for r_idx, row in enumerate(body, start=1):
        values = [x.strip() for x in row.strip().strip("|").split("|")]
        for c_idx, text in enumerate(values):
            cell = table.rows[r_idx].cells[c_idx]
            p = cell.paragraphs[0]
            p.style = "Table Text"
            p.clear()
            add_inline(p, text, base_size=9, base_color=INK)

    after = doc.add_paragraph(style="Source Note")
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    return table


def render_markdown(doc, path, part_title, skip_title=True):
    p = doc.add_paragraph(part_title, style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    first_h1_skipped = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                add_markdown_table(doc, table_lines)
            continue
        if line.startswith("# "):
            if skip_title and not first_h1_skipped:
                first_h1_skipped = True
            else:
                doc.add_paragraph(line[2:].strip(), style="Heading 1")
            i += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_paragraph(line[5:].strip(), style="Heading 3")
            i += 1
            continue
        if line.startswith("> "):
            p = doc.add_paragraph(style="Source Note")
            add_inline(p, line[2:], base_size=9, base_color=MUTED)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph()
            apply_numbering(p, doc._numbering_ids["decimal"])
            add_inline(p, text)
            i += 1
            continue
        if line.startswith("- [ ] "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            add_inline(p, "□ " + line[6:])
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph()
            apply_numbering(p, doc._numbering_ids["bullet"])
            add_inline(p, line[2:])
            i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("#", "|", "> ", "- ", "- [ ] ")) or re.match(r"^\d+\.\s+", nxt):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))


def add_document_properties(doc):
    props = doc.core_properties
    props.title = "AI 市场工作台：二手证据验证与市场测试执行包"
    props.subject = "公开调研、MVP 判断、真实 brief 回放与付费行为测试"
    props.author = "MarketOps AI Workbench Project"
    props.keywords = "AI, 市场营销, 活动策划, brief, 市场验证, GEO, 工作台"
    props.comments = "审阅版 v0.1，2026-08-08"


def audit_structure(doc):
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.50
    assert round(section.page_height.inches, 2) == 11.00
    assert all(
        round(value.inches, 2) == 1.00
        for value in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin)
    )
    assert round(section.header_distance.inches, 3) == 0.492
    assert round(section.footer_distance.inches, 3) == 0.492
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        assert tbl_w is not None and tbl_w.get(qn("w:w")) == "9360"
        grid_widths = [int(x.get(qn("w:w"))) for x in table._tbl.tblGrid]
        assert sum(grid_widths) == 9360


def build():
    doc = Document()
    configure_styles(doc)
    setup_page(doc)
    doc._numbering_ids = add_numbering_definitions(doc)
    add_document_properties(doc)
    add_masthead(doc)
    add_contents(doc)
    render_markdown(doc, EVIDENCE_MD, "Part I  二手证据与产品判断")
    doc.add_page_break()
    render_markdown(doc, PLAYBOOK_MD, "Part II  市场验证执行包")
    audit_structure(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
